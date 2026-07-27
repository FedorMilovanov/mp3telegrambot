from __future__ import annotations

import asyncio
import html
import os
import re
from pathlib import Path
from typing import Any

from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import ApplicationHandlerStop, ContextTypes

from core.database import ADMIN_IDS, WHITELIST_IDS
from core.dub_projects import (
    DubProjectError,
    assert_project_owner,
    attach_approved_translation,
    attach_source_file,
    cancel_project,
    create_project,
    extract_project_id,
    load_project,
    manifest_path,
    project_dir,
    project_marker,
)
from core.utils import extract_media_url, is_media_url
from pipelines.dubbing.preflight import run_project_preflight


_ALLOWED_TRANSLATION_SUFFIXES = {".txt", ".md", ".docx"}
_ALLOWED_VIDEO_SUFFIXES = {".mp4", ".mov", ".mkv", ".webm", ".m4v", ".avi"}
_SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9А-Яа-яЁё._ -]+")


def _production_access(user_id: int) -> bool:
    return int(user_id) in set(ADMIN_IDS or []) or int(user_id) in set(WHITELIST_IDS or [])


def _admins() -> set[int]:
    return {int(value) for value in (ADMIN_IDS or [])}


def _safe_filename(value: str | None, fallback: str) -> str:
    name = Path(str(value or "")).name.strip() or fallback
    name = _SAFE_FILENAME_RE.sub("_", name).strip(" .")
    return (name or fallback)[:180]


def _project_keyboard(project_id: str, *, translation_ready: bool) -> InlineKeyboardMarkup:
    rows: list[list[InlineKeyboardButton]] = []
    if translation_ready:
        rows.append([InlineKeyboardButton("🔎 Проверить готовность", callback_data=f"dub:preflight:{project_id}")])
        rows.append([InlineKeyboardButton("📝 Заменить перевод", callback_data=f"dub:replace:{project_id}")])
    rows.append(
        [
            InlineKeyboardButton("📋 Manifest", callback_data=f"dub:manifest:{project_id}"),
            InlineKeyboardButton("✖️ Отменить", callback_data=f"dub:cancel:{project_id}"),
        ]
    )
    return InlineKeyboardMarkup(rows)


def _source_from_message(message: Any) -> tuple[dict[str, Any], Any] | None:
    media = getattr(message, "video", None)
    if media is not None:
        return (
            {
                "kind": "telegram_file",
                "file_id": media.file_id,
                "file_unique_id": media.file_unique_id,
                "filename": _safe_filename(getattr(media, "file_name", None), "source.mp4"),
                "mime_type": getattr(media, "mime_type", None) or "video/mp4",
                "file_size": int(getattr(media, "file_size", 0) or 0),
            },
            media,
        )
    document = getattr(message, "document", None)
    if document is None:
        return None
    suffix = Path(str(document.file_name or "")).suffix.lower()
    mime_type = str(document.mime_type or "").lower()
    if suffix not in _ALLOWED_VIDEO_SUFFIXES and not mime_type.startswith("video/"):
        return None
    return (
        {
            "kind": "telegram_file",
            "file_id": document.file_id,
            "file_unique_id": document.file_unique_id,
            "filename": _safe_filename(document.file_name, "source.mp4"),
            "mime_type": document.mime_type or "application/octet-stream",
            "file_size": int(document.file_size or 0),
        },
        document,
    )


def _reply_text(message: Any) -> str:
    return str(getattr(message, "text", None) or getattr(message, "caption", None) or "")


def _translation_prompt(project_id: str, *, replacement: bool = False) -> str:
    action = "новый утверждённый перевод" if replacement else "уже проверенный и окончательно утверждённый перевод"
    return (
        "🎬 <b>Production-проект создан</b>\n\n"
        f"Ответьте <b>на это сообщение</b> и пришлите {action}:\n"
        "• обычным текстом;\n"
        "• файлом <code>.txt</code>, <code>.md</code> или <code>.docx</code>.\n\n"
        "Бот не будет переводить, сокращать или литературно переписывать текст. "
        "Он сохранит утверждённую версию, проверит её целостность и подготовит производственный preflight.\n\n"
        f"<code>{project_marker(project_id)}</code>"
    )


def _translation_summary(manifest: dict[str, Any]) -> str:
    project_id = str(manifest["project_id"])
    translation = manifest.get("translation") or {}
    source = manifest.get("source") or {}
    source_value = source.get("url") or source.get("filename") or source.get("kind")
    return (
        "✅ <b>Утверждённый перевод принят</b>\n\n"
        f"Проект: <code>{html.escape(project_id)}</code>\n"
        f"Источник: <code>{html.escape(str(source_value or 'не указан'))}</code>\n"
        f"Ревизия перевода: <b>{int(translation.get('revision') or 1)}</b>\n"
        f"Символов: <b>{int(translation.get('character_count') or 0)}</b>\n"
        f"Слов: <b>{int(translation.get('word_count') or 0)}</b>\n"
        f"Редакционных блоков: <b>{int(translation.get('unit_count') or 0)}</b>\n"
        f"SHA-256: <code>{html.escape(str(translation.get('sha256') or '')[:16])}…</code>\n\n"
        "🔒 Перевод заблокирован как редакционно утверждённый. Изменение файла снимет готовность production-проекта."
    )


def _preflight_text(report: dict[str, Any]) -> str:
    ok = bool(report.get("ok"))
    duration = float(report.get("duration_seconds") or 0)
    profile = str(report.get("profile") or "unknown")
    subtitles = report.get("subtitles") or {}
    disk = report.get("disk") or {}
    lines = [
        ("✅ <b>Preflight пройден</b>" if ok else "❌ <b>Preflight не пройден</b>"),
        "",
        f"Профиль: <code>{html.escape(profile)}</code>",
        f"Длительность: <b>{duration:.3f} сек.</b>",
        f"Прожиг субтитров: <b>{'да' if subtitles.get('hardsub') else 'нет'}</b>",
        "Отдельный SRT: <b>да</b>",
        "Надписи в кадре: <b>не переводятся</b>",
        "Движок: <b>VoxCPM2 / CPU</b>",
        "Скрытый TTS fallback: <b>запрещён</b>",
        f"Диск: {float(disk.get('free_gib') or 0):.1f} ГБ свободно / около {float(disk.get('estimated_required_gib') or 0):.1f} ГБ требуется",
    ]
    blocking = report.get("blocking_errors") or []
    warnings = report.get("warnings") or []
    if blocking:
        lines.extend(["", "<b>Блокирующие ошибки:</b>"])
        lines.extend(f"• {html.escape(str(item))}" for item in blocking[:12])
    if warnings:
        lines.extend(["", "<b>Предупреждения:</b>"])
        lines.extend(f"• {html.escape(str(item))}" for item in warnings[:8])
    if ok:
        lines.extend(
            [
                "",
                "Статический и исходный preflight завершён. Следующий production-этап — автоматическое получение стенограммы, сопоставление утверждённого перевода, голосовые референсы и сегментированный синтез.",
            ]
        )
    return "\n".join(lines)[:3900]


async def _download_source(context: ContextTypes.DEFAULT_TYPE, project_id: str, source: dict[str, Any]) -> Path:
    max_mb = int(os.getenv("DUB_MAX_SOURCE_MB", "2000"))
    size = int(source.get("file_size") or 0)
    if size and size > max_mb * 1024 * 1024:
        raise DubProjectError(f"Исходный файл больше лимита {max_mb} МБ.")
    filename = _safe_filename(source.get("filename"), "source.mp4")
    suffix = Path(filename).suffix.lower() or ".mp4"
    target = project_dir(project_id) / "source" / f"source{suffix}"
    tg_file = await context.bot.get_file(str(source["file_id"]))
    await tg_file.download_to_drive(custom_path=target)
    if not target.is_file() or target.stat().st_size <= 0:
        raise DubProjectError("Telegram не сохранил исходный видеофайл.")
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, attach_source_file, project_id, target)
    return target


async def dub_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    user = update.effective_user
    if message is None or user is None:
        return
    if not _production_access(user.id):
        await message.reply_text("⛔ Production-дубляж доступен только владельцу/VIP.")
        return

    source: dict[str, Any] | None = None
    media_descriptor: Any = None
    args_text = " ".join(context.args or []).strip()
    if args_text and is_media_url(args_text):
        url = extract_media_url(args_text)
        if url:
            source = {"kind": "url", "url": url if url.startswith("http") else "https://" + url}
    reply = message.reply_to_message
    if source is None and reply is not None:
        media_source = _source_from_message(reply)
        if media_source is not None:
            source, media_descriptor = media_source
        else:
            candidate = _reply_text(reply)
            if is_media_url(candidate):
                url = extract_media_url(candidate)
                if url:
                    source = {"kind": "url", "url": url if url.startswith("http") else "https://" + url}

    if source is None:
        await message.reply_text(
            "🎬 <b>Production-дубляж VoxCPM2</b>\n\n"
            "Использование:\n"
            "• <code>/dub https://youtube.com/...</code>\n"
            "• ответьте командой <code>/dub</code> на сообщение с исходным видеофайлом.\n\n"
            "После создания проекта бот попросит уже проверенный русский перевод.",
            parse_mode="HTML",
        )
        return

    loop = asyncio.get_running_loop()
    try:
        manifest = await loop.run_in_executor(None, lambda: create_project(owner_user_id=user.id, source=source))
        project_id = str(manifest["project_id"])
        if source["kind"] == "telegram_file":
            status = await message.reply_text("⬇️ Сохраняю исходный видеофайл в production-проект…")
            try:
                await _download_source(context, project_id, source)
            finally:
                try:
                    await status.delete()
                except Exception:
                    pass
    except DubProjectError as exc:
        await message.reply_text(f"❌ {html.escape(str(exc))}", parse_mode="HTML")
        return
    except Exception as exc:
        await message.reply_text(f"❌ Не удалось создать production-проект: {html.escape(str(exc)[:300])}", parse_mode="HTML")
        return

    await message.reply_text(
        _translation_prompt(project_id),
        parse_mode="HTML",
        reply_markup=_project_keyboard(project_id, translation_ready=False),
    )


def _marker_from_reply(message: Any) -> str | None:
    reply = getattr(message, "reply_to_message", None)
    if reply is None:
        return None
    return extract_project_id(_reply_text(reply))


async def _read_translation_document(message: Any, context: ContextTypes.DEFAULT_TYPE, project_id: str) -> tuple[str, str]:
    document = message.document
    suffix = Path(str(document.file_name or "")).suffix.lower()
    if suffix not in _ALLOWED_TRANSLATION_SUFFIXES:
        raise DubProjectError("Перевод принимается только как TXT, MD или DOCX.")
    max_mb = int(os.getenv("DUB_MAX_TRANSLATION_MB", "10"))
    if int(document.file_size or 0) > max_mb * 1024 * 1024:
        raise DubProjectError(f"Файл перевода больше лимита {max_mb} МБ.")
    filename = _safe_filename(document.file_name, "translation" + suffix)
    target = project_dir(project_id) / "incoming" / filename
    tg_file = await context.bot.get_file(document.file_id)
    await tg_file.download_to_drive(custom_path=target)

    def read() -> str:
        if suffix in {".txt", ".md"}:
            raw = target.read_bytes()
            for encoding in ("utf-8-sig", "utf-16", "cp1251"):
                try:
                    return raw.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise DubProjectError("Не удалось определить кодировку текстового файла.")
        try:
            from docx import Document
        except ImportError as exc:
            raise DubProjectError("Для DOCX установите зависимости: pip install -r requirements.txt") from exc
        doc = Document(str(target))
        blocks: list[str] = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    blocks.append(" | ".join(values))
        return "\n\n".join(blocks)

    text = await asyncio.get_running_loop().run_in_executor(None, read)
    return text, filename


async def _attach_translation_from_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> bool:
    message = update.effective_message
    user = update.effective_user
    if message is None or user is None:
        return False
    project_id = _marker_from_reply(message)
    if project_id is None:
        return False
    if not _production_access(user.id):
        await message.reply_text("⛔ Нет доступа к production-дубляжу.")
        raise ApplicationHandlerStop

    try:
        manifest = load_project(project_id)
        assert_project_owner(manifest, user.id, admin_ids=_admins())
        if message.document is not None:
            text, filename = await _read_translation_document(message, context, project_id)
        else:
            text = str(message.text or "")
            filename = "telegram-message"
        manifest = await asyncio.get_running_loop().run_in_executor(
            None,
            lambda: attach_approved_translation(
                project_id,
                text=text,
                approved_by_user_id=user.id,
                original_filename=filename,
            ),
        )
    except DubProjectError as exc:
        await message.reply_text(f"❌ {html.escape(str(exc))}", parse_mode="HTML")
        raise ApplicationHandlerStop
    except Exception as exc:
        await message.reply_text(f"❌ Не удалось принять перевод: {html.escape(str(exc)[:300])}", parse_mode="HTML")
        raise ApplicationHandlerStop

    await message.reply_text(
        _translation_summary(manifest),
        parse_mode="HTML",
        reply_markup=_project_keyboard(project_id, translation_ready=True),
    )
    raise ApplicationHandlerStop


async def handle_dub_translation_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await _attach_translation_from_message(update, context)


async def handle_dub_translation_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await _attach_translation_from_message(update, context)


async def handle_dub_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    user = update.effective_user
    if query is None or user is None:
        return
    data = str(query.data or "")
    parts = data.split(":", 2)
    if len(parts) != 3 or parts[0] != "dub":
        await query.answer("Некорректная команда.", show_alert=True)
        return
    action, project_id = parts[1], parts[2]
    try:
        manifest = load_project(project_id)
        assert_project_owner(manifest, user.id, admin_ids=_admins())
    except DubProjectError as exc:
        await query.answer(str(exc), show_alert=True)
        return

    if action == "replace":
        await query.answer()
        await query.message.reply_text(
            _translation_prompt(project_id, replacement=True),
            parse_mode="HTML",
            reply_markup=_project_keyboard(project_id, translation_ready=True),
        )
        return

    if action == "manifest":
        await query.answer()
        path = manifest_path(project_id)
        with path.open("rb") as stream:
            await query.message.reply_document(
                document=stream,
                filename=f"{project_id}.manifest.json",
                caption=f"📋 {project_marker(project_id)}",
                write_timeout=180,
                read_timeout=180,
                connect_timeout=60,
            )
        return

    if action == "cancel":
        await query.answer("Проект отменён")
        await asyncio.get_running_loop().run_in_executor(
            None, lambda: cancel_project(project_id, cancelled_by_user_id=user.id)
        )
        try:
            await query.edit_message_text(
                f"✖️ Production-проект отменён.\n\n<code>{project_marker(project_id)}</code>",
                parse_mode="HTML",
            )
        except Exception:
            pass
        return

    if action == "preflight":
        await query.answer("Запускаю preflight…")
        status = await query.message.reply_text("🔎 Проверяю исходник, перевод, VoxCPM2 CPU, FFmpeg и место на диске…")
        try:
            report = await asyncio.get_running_loop().run_in_executor(None, run_project_preflight, project_id)
            await status.edit_text(_preflight_text(report), parse_mode="HTML")
        except Exception as exc:
            await status.edit_text(f"❌ Preflight завершился ошибкой: {html.escape(str(exc)[:500])}", parse_mode="HTML")
        return

    await query.answer("Неизвестное действие.", show_alert=True)
